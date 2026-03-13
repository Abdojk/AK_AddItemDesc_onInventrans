"""extractor.py — Extract text content from .docx files for SOW01 analysis."""

import re
import sys
from pathlib import Path

import docx

MAX_CHARS = 12_000


def extract_text(docx_path: Path, verbose: bool = False) -> str:
    """Extract all text from a .docx file and return as a single string.

    Extracts paragraphs in document order, then table cells with positional
    prefixes. Detects inline images and warns if present. Truncates output
    to MAX_CHARS characters.
    """
    try:
        document = docx.Document(docx_path)
    except Exception as e:
        print(f"Error: Could not open '{docx_path}': {e}", file=sys.stderr)
        sys.exit(1)

    lines: list[str] = []

    # Extract paragraphs in document order
    for para in document.paragraphs:
        lines.append(para.text)

    # Extract table cells with positional prefixes
    for table_idx, table in enumerate(document.tables, start=1):
        lines.append("")  # blank line before each table
        for row_idx, row in enumerate(table.rows, start=1):
            for col_idx, cell in enumerate(row.cells, start=1):
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(
                        f"[Table {table_idx} | Row {row_idx} | Col {col_idx}] {cell_text}"
                    )

    # Detect inline shapes (images)
    image_count = len(document.inline_shapes)
    if image_count > 0:
        lines.append("")
        lines.append(
            f"\u26a0 WARNING: {image_count} image(s) detected. Image content could not "
            "be extracted. If scope information is embedded in screenshots, the "
            "analysis may be incomplete. Consider pasting screenshot content as text."
        )

    # Join and collapse consecutive blank lines
    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    # Truncate if necessary
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]
        truncation_warning = (
            "\u26a0 Input truncated to 12,000 characters. Long SOW documents may "
            "produce incomplete analysis."
        )
        print(truncation_warning, file=sys.stderr)
        text += f"\n\n{truncation_warning}"

    if verbose:
        print(f"Extracted {len(text)} characters from '{docx_path.name}'.", file=sys.stderr)

    return text
