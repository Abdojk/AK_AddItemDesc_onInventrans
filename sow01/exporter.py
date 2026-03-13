"""exporter.py — Export SOW01 assessment to markdown, DOCX, or stdout."""

import datetime
import sys
from pathlib import Path

import docx as python_docx
from docx.shared import Pt, Inches


def _build_output_stem(source_filename: str) -> str:
    """Build the output filename stem: SOW01_[original]_[YYYYMMDD]."""
    date_str = datetime.date.today().strftime("%Y%m%d")
    return f"SOW01_{source_filename}_{date_str}"


def _export_markdown(
    assessment: str,
    source_filename: str,
    output_dir: Path,
    stem: str,
    verbose: bool = False,
) -> Path:
    """Write assessment to a markdown file with a header block."""
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    header = (
        f"# SOW01 Assessment Report\n\n"
        f"**Source document:** {source_filename}.docx  \n"
        f"**Date generated:** {date_str}  \n"
        f"**Tool:** SOW01 — D365 F&O Statement of Work Reviewer\n\n"
        f"---\n\n"
    )
    content = header + assessment + "\n"
    output_path = output_dir / f"{stem}.md"
    output_path.write_text(content, encoding="utf-8")

    if verbose:
        print(f"Markdown saved to: {output_path}", file=sys.stderr)

    return output_path


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    """Parse pipe-delimited markdown table lines into a list of rows.

    Skips separator rows (containing only dashes and pipes).
    """
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        # Skip separator rows like |---|---|---|
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        if all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)
    return rows


def _export_docx(
    assessment: str,
    source_filename: str,
    output_dir: Path,
    stem: str,
    verbose: bool = False,
) -> Path:
    """Write assessment to a formatted Word document."""
    doc = python_docx.Document()

    # Title
    doc.add_heading("SOW01 Assessment Report", level=0)

    # Metadata
    date_str = datetime.date.today().strftime("%Y-%m-%d")
    meta = doc.add_paragraph()
    meta.add_run("Source document: ").bold = True
    meta.add_run(f"{source_filename}.docx")
    meta = doc.add_paragraph()
    meta.add_run("Date generated: ").bold = True
    meta.add_run(date_str)
    meta = doc.add_paragraph()
    meta.add_run("Tool: ").bold = True
    meta.add_run("SOW01 — D365 F&O Statement of Work Reviewer")

    doc.add_paragraph("")  # spacer

    # Parse assessment line by line
    assessment_lines = assessment.split("\n")
    table_lines: list[str] = []
    in_table = False

    for line in assessment_lines:
        stripped = line.strip()

        # Detect table rows
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            continue

        # If we were in a table and hit a non-table line, flush the table
        if in_table:
            _add_table_to_doc(doc, table_lines)
            table_lines = []
            in_table = False

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped:
            # Replace <br> with newlines for Word
            doc.add_paragraph(stripped.replace("<br>", "\n"))

    # Flush any remaining table
    if table_lines:
        _add_table_to_doc(doc, table_lines)

    output_path = output_dir / f"{stem}.docx"
    doc.save(str(output_path))

    if verbose:
        print(f"DOCX saved to: {output_path}", file=sys.stderr)

    return output_path


def _add_table_to_doc(doc: python_docx.Document, table_lines: list[str]) -> None:
    """Parse markdown table lines and add a formatted table to the Word document."""
    rows = _parse_markdown_table(table_lines)
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols, style="Light Grid Accent 1")

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                # Replace <br> tags with newlines
                cell.text = cell_text.replace("<br>", "\n")
                # Bold the header row
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def export_results(
    assessment: str,
    source_filename: str,
    fmt: str,
    output_dir: Path,
    verbose: bool = False,
) -> None:
    """Export assessment results based on the chosen format.

    Args:
        assessment: The raw assessment text from the API.
        source_filename: Original input filename (without extension).
        fmt: Output format — "md", "docx", "both", or "none".
        output_dir: Directory to write output files.
        verbose: Print diagnostic info to stderr.
    """
    # Always print assessment to stdout
    print(assessment)

    if fmt == "none":
        return

    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = _build_output_stem(source_filename)

    if fmt in ("md", "both"):
        _export_markdown(assessment, source_filename, output_dir, stem, verbose)

    if fmt in ("docx", "both"):
        _export_docx(assessment, source_filename, output_dir, stem, verbose)
